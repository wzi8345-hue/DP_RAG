from pathlib import Path
import textwrap
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "assets_function_doc"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "DP-RAG功能描述文档.docx"
FONT_PATHS = [
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
]
FONT = next((p for p in FONT_PATHS if Path(p).exists()), None)

COLORS = {
    "navy": "12335B",
    "blue": "2563EB",
    "sky": "EAF3FF",
    "green": "0F766E",
    "mint": "E6FFFB",
    "orange": "EA580C",
    "cream": "FFF7ED",
    "purple": "7C3AED",
    "lavender": "F3E8FF",
    "gray": "64748B",
    "line": "CBD5E1",
    "text": "0F172A",
    "white": "FFFFFF",
}

def font(size, bold=False):
    return ImageFont.truetype(FONT, size=size) if FONT else ImageFont.load_default()

def hex_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0,2,4))

def rounded(draw, box, fill, outline="CBD5E1", width=3, radius=22):
    draw.rounded_rectangle(box, radius=radius, fill=hex_rgb(fill), outline=hex_rgb(outline), width=width)

def wrap_label(draw, text, max_width, fnt):
    lines = []
    for raw in str(text).split("\n"):
        current = ""
        for ch in raw:
            trial = current + ch
            if draw.textbbox((0,0), trial, font=fnt)[2] <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines or [""]

def centered_text(draw, box, text, fnt, fill="0F172A", line_gap=8):
    x1, y1, x2, y2 = box
    maxw = x2 - x1 - 28
    lines = wrap_label(draw, text, maxw, fnt)
    heights = [draw.textbbox((0,0), line, font=fnt)[3] - draw.textbbox((0,0), line, font=fnt)[1] for line in lines]
    total_h = sum(heights) + line_gap * (len(lines)-1)
    y = y1 + (y2-y1-total_h)/2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0,0), line, font=fnt)
        tw = bbox[2]-bbox[0]
        draw.text((x1 + (x2-x1-tw)/2, y), line, font=fnt, fill=hex_rgb(fill))
        y += h + line_gap

def arrow(draw, start, end, color="64748B", width=5):
    draw.line([start, end], fill=hex_rgb(color), width=width)
    import math
    sx, sy = start; ex, ey = end
    ang = math.atan2(ey-sy, ex-sx)
    size = 16
    pts = [
        (ex, ey),
        (ex - size*math.cos(ang-0.45), ey - size*math.sin(ang-0.45)),
        (ex - size*math.cos(ang+0.45), ey - size*math.sin(ang+0.45)),
    ]
    draw.polygon(pts, fill=hex_rgb(color))

def title(draw, text, subtitle=None):
    draw.text((50, 34), text, font=font(34, True), fill=hex_rgb(COLORS["navy"]))
    if subtitle:
        draw.text((52, 82), subtitle, font=font(20), fill=hex_rgb(COLORS["gray"]))

def save_diagram(name, w, h, drawer):
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    drawer(d, w, h)
    path = ASSET_DIR / f"{name}.png"
    img.save(path, quality=95)
    return path

def box(draw, xy, text, fill="EAF3FF", outline="2563EB", txt="0F172A", fs=23):
    rounded(draw, xy, fill, outline)
    centered_text(draw, xy, text, font(fs), txt)

# Diagram 1: architecture
save_diagram("01_overall_architecture", 1500, 720, lambda d,w,h: (
    title(d, "总体架构", "从 PDF 文献到带引用的科研问答"),
    [box(d, (60+i*220, 180, 220+i*220, 285), t, f, o, fs=22) for i,(t,f,o) in enumerate([
        ("PDF\n文献", "FFF7ED", "EA580C"), ("解析\nMinerU/UniParser", "EAF3FF", "2563EB"), ("知识\n分块", "E6FFFB", "0F766E"), ("Embedding\n向量化", "F3E8FF", "7C3AED"), ("Milvus\n向量库", "EAF3FF", "2563EB")])],
    [arrow(d, (220+i*220,232), (60+(i+1)*220,232)) for i in range(4)],
    box(d, (60, 470, 220, 575), "用户\n问题", "FFF7ED", "EA580C"),
    box(d, (330, 470, 520, 575), "查询路由\n与检索", "EAF3FF", "2563EB"),
    box(d, (650, 470, 845, 575), "上下文\n构建", "E6FFFB", "0F766E"),
    box(d, (970, 470, 1160, 575), "LLM\n生成", "F3E8FF", "7C3AED"),
    box(d, (1285, 470, 1450, 575), "答案 +\n引用来源", "EAF3FF", "2563EB"),
    arrow(d, (220,522), (330,522)), arrow(d, (520,522), (650,522)), arrow(d, (845,522), (970,522)), arrow(d, (1160,522), (1285,522)),
    arrow(d, (940,285), (740,470), "94A3B8", 4)
))

# Diagram 2: ingest flow
save_diagram("02_ingest_flow", 1500, 560, lambda d,w,h: (
    title(d, "功能一：文献解析与知识灌入", "解析、分块、向量化、入库一体化"),
    [box(d, (70+i*235, 210, 250+i*235, 330), t, f, o, fs=22) for i,(t,f,o) in enumerate([
        ("上传 PDF\n或选择目录", "FFF7ED", "EA580C"), ("解析文献\n结构化内容", "EAF3FF", "2563EB"), ("生成\n知识块", "E6FFFB", "0F766E"), ("生成\n向量", "F3E8FF", "7C3AED"), ("写入\nMilvus", "EAF3FF", "2563EB"), ("返回任务\n进度结果", "E6FFFB", "0F766E")])],
    [arrow(d, (250+i*235,270), (70+(i+1)*235,270)) for i in range(5)],
    d.text((90, 390), "支持：单篇/批量导入、重建/追加、跳过已存在、加载已有向量 JSON。", font=font(24), fill=hex_rgb(COLORS["gray"]))
))

# Diagram 3: query flow
save_diagram("03_query_flow", 1500, 760, lambda d,w,h: (
    title(d, "功能二：智能问答", "按查询模式检索证据，并生成带来源回答"),
    box(d, (650, 135, 850, 230), "用户输入\n科研问题", "FFF7ED", "EA580C", fs=24),
    box(d, (650, 310, 850, 405), "选择\n查询模式", "EAF3FF", "2563EB", fs=24),
    arrow(d, (750,230), (750,310)),
    box(d, (170, 500, 390, 610), "简单模式\nVector / Metadata / Hybrid", "EAF3FF", "2563EB", fs=21),
    box(d, (640, 500, 860, 610), "Agentic RAG\n多路径并行检索", "E6FFFB", "0F766E", fs=21),
    box(d, (1110, 500, 1330, 610), "专业研究模式\n多轮递进检索", "F3E8FF", "7C3AED", fs=21),
    arrow(d, (650,405), (280,500)), arrow(d, (750,405), (750,500)), arrow(d, (850,405), (1220,500)),
    box(d, (520, 650, 980, 725), "融合证据 → 构建上下文 → LLM 生成答案 → 展示引用来源", "FFF7ED", "EA580C", fs=24),
    arrow(d, (390,610), (590,650)), arrow(d, (750,610), (750,650)), arrow(d, (1110,610), (910,650))
))

# Diagram 4 retrieval
save_diagram("04_retrieval_strategy", 1500, 760, lambda d,w,h: (
    title(d, "功能三：检索增强与专业研究", "多路召回、上下文扩展、重排与质量门控"),
    box(d, (620, 125, 880, 220), "查询路由器", "FFF7ED", "EA580C", fs=25),
    [box(d, (x, 330, x+230, 430), t, f, o, fs=21) for x,t,f,o in [
        (90,"元数据检索\n标题/年份/实体", "EAF3FF", "2563EB"),
        (390,"向量检索\n语义相似度", "E6FFFB", "0F766E"),
        (690,"结构化检索\n章节/摘要/表格", "F3E8FF", "7C3AED"),
        (990,"专家技能\n策略与提示词", "FFF7ED", "EA580C"),
        (1240,"邻居补充\n上下文扩展", "EAF3FF", "2563EB")]],
    [arrow(d, (750,220), (x+115,330)) for x in [90,390,690,990,1240]],
    box(d, (420, 560, 680, 650), "结果融合", "EAF3FF", "2563EB", fs=24),
    box(d, (810, 560, 1070, 650), "Reranker\n重排诊断", "E6FFFB", "0F766E", fs=24),
    box(d, (1200, 560, 1430, 650), "最终上下文", "F3E8FF", "7C3AED", fs=24),
    [arrow(d, (x+115,430), (550,560)) for x in [90,390,690,990,1240]],
    arrow(d, (680,605), (810,605)), arrow(d, (1070,605), (1200,605))
))

# Diagram 5 mgmt
save_diagram("05_management_flow", 1500, 690, lambda d,w,h: (
    title(d, "功能四：知识库、技能与系统管理", "前端统一管理文献库、专家技能和运行状态"),
    box(d, (80, 170, 270, 270), "前端界面", "FFF7ED", "EA580C", fs=24),
    [box(d, (470, y, 720, y+90), t, f, o, fs=22) for y,t,f,o in [
        (100,"知识库管理\ncollections / tasks", "EAF3FF", "2563EB"),
        (260,"专家技能管理\nskills / template", "E6FFFB", "0F766E"),
        (420,"系统状态\nhealth / stats", "F3E8FF", "7C3AED")]],
    [arrow(d, (270,220), (470,y+45)) for y in [100,260,420]],
    [box(d, (950, y, 1320, y+90), t, f, o, fs=22) for y,t,f,o in [
        (100,"查看/删除 Collection\n导入任务进度追踪", "EAF3FF", "2563EB"),
        (260,"编辑触发条件、检索策略\n和合成提示词", "E6FFFB", "0F766E"),
        (420,"检查 Milvus、LLM、Embedding\nReranker 与日志", "F3E8FF", "7C3AED")]],
    [arrow(d, (720,y+45), (950,y+45)) for y in [100,260,420]],
    d.text((90, 590), "辅助能力：会话日志、文档摘要、RAGAS 评测、合成问答数据生成。", font=font(24), fill=hex_rgb(COLORS["gray"]))
))

# Diagram 6 frontend map
save_diagram("06_frontend_map", 1500, 700, lambda d,w,h: (
    title(d, "功能五：前端功能地图", "React + Vite 提供可视化科研问答工作台"),
    box(d, (610, 120, 890, 220), "DP-RAG 工作台", "FFF7ED", "EA580C", fs=26),
    [box(d, xy, text, fill, outline, fs=22) for xy,text,fill,outline in [
        ((70,350,285,455), "智能问答\n答案/引用/历史", "EAF3FF", "2563EB"),
        ((330,350,545,455), "知识库\n上传/导入/任务", "E6FFFB", "0F766E"),
        ((590,350,805,455), "专家技能\n编辑/删除", "F3E8FF", "7C3AED"),
        ((850,350,1065,455), "系统状态\n依赖/统计", "FFF7ED", "EA580C"),
        ((1110,350,1325,455), "检索日志\n会话链路", "EAF3FF", "2563EB")]],
    [arrow(d, (750,220), (x,350)) for x in [177,437,697,957,1217]],
    d.text((90, 560), "侧栏统一进入各模块；设置弹窗用于配置 API 地址和鉴权信息。", font=font(24), fill=hex_rgb(COLORS["gray"]))
))

# DOCX helpers

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, color="0F172A"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
        run.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    return p

def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
    run.font.size = Pt(10.5)
    p.paragraph_format.line_spacing = 1.25
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
        run.font.size = Pt(10.5)

def add_diagram(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(6.65))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor.from_string(COLORS["gray"])

# Build document
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.65)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.65)
sec.right_margin = Inches(0.65)

styles = doc.styles
styles['Normal'].font.name = 'PingFang SC'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
styles['Normal'].font.size = Pt(10.5)

# Cover
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("DP-RAG 功能描述文档")
title_run.font.name = "PingFang SC"
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor.from_string(COLORS["navy"])
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("面向科研文献的端到端 RAG 知识库与智能问答系统")
r.font.name = "PingFang SC"
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string(COLORS["gray"])
add_para(doc, "本文档以流程图为主，简要说明项目的核心功能、业务流程、前后端模块和典型使用场景。")
add_diagram(doc, ASSET_DIR / "01_overall_architecture.png", "图 1：DP-RAG 总体架构")

add_heading(doc, "1. 项目概述", 1)
add_para(doc, "DP-RAG 是一个面向科研文献的检索增强生成系统，覆盖 PDF 解析、知识分块、向量化入库、混合检索、智能体式检索和 LLM 答案生成。系统由后端 RAG 流水线、FastAPI 服务、React 前端、评测工具和合成问答数据工具组成。")
add_bullets(doc, [
    "目标用户：科研人员、文献分析人员、需要基于论文证据进行问答和综述的团队。",
    "核心价值：把科研 PDF 转换为可检索、可引用、可追踪的知识库。",
    "输出形式：答案、命中文献片段、引用来源、上下文、执行日志和统计信息。",
])

add_heading(doc, "2. 文献解析与知识灌入", 1)
add_diagram(doc, ASSET_DIR / "02_ingest_flow.png", "图 2：文献解析与知识灌入流程")
add_para(doc, "该功能负责把原始 PDF 文献转化为 Milvus 中的知识向量记录，是整个 RAG 系统的基础数据入口。")
add_bullets(doc, [
    "支持单篇 PDF、目录批量导入和前端上传导入。",
    "支持 MinerU / UniParser 等解析后端，生成结构化文档内容。",
    "自动完成知识分块、Embedding 向量化和 Milvus 写入。",
    "支持重建知识库、追加导入、跳过已存在文档和加载已有向量 JSON。",
])

add_heading(doc, "3. 智能问答", 1)
add_diagram(doc, ASSET_DIR / "03_query_flow.png", "图 3：智能问答流程")
add_para(doc, "用户可以基于已入库的科研文献进行自然语言提问，系统会检索相关证据、构建上下文，并调用 LLM 生成带引用的答案。")
add_bullets(doc, [
    "简单模式：适合快速查询，可指定 vector、metadata 或 hybrid 检索。",
    "Agentic RAG：使用路由和多路径检索提升复杂问题的召回质量。",
    "专业研究模式：面向综述类问题，支持多轮递进式文献检索和证据汇总。",
    "前端支持流式回答、历史对话、引用来源面板和命中文档查看。",
])

add_heading(doc, "4. 检索增强与专业研究", 1)
add_diagram(doc, ASSET_DIR / "04_retrieval_strategy.png", "图 4：检索增强策略关系")
add_para(doc, "项目内置多种检索增强机制，用于提高科研场景下的召回覆盖、证据质量和答案可信度。")
add_bullets(doc, [
    "元数据检索：利用标题、年份、文档名、实体等信息缩小检索范围。",
    "向量检索：通过 Embedding 相似度召回语义相关片段。",
    "结构化检索：围绕章节、摘要、表格等结构信息补充证据。",
    "邻居扩展：对命中片段扩展上下文，避免孤立片段造成理解偏差。",
    "Reranker 与质量门控：对候选证据重排、诊断并筛选更可靠上下文。",
])

add_heading(doc, "5. 知识库、技能与系统管理", 1)
add_diagram(doc, ASSET_DIR / "05_management_flow.png", "图 5：知识库、专家技能和系统状态管理流程")
add_para(doc, "除核心问答外，系统还提供知识库维护、专家技能配置、运行状态检查和检索日志查看能力。")
add_bullets(doc, [
    "知识库管理：查看/删除 collection，跟踪导入任务，查看文档摘要和统计信息。",
    "专家技能管理：配置触发条件、检索优先级、充分性标准、保护规则和合成提示词。",
    "系统状态：检查 Milvus、LLM、Embedding、Reranker、Reflection 等依赖是否可用。",
    "检索日志：按会话查看检索链路，便于问题定位和效果调试。",
])

add_heading(doc, "6. 前端功能地图", 1)
add_diagram(doc, ASSET_DIR / "06_frontend_map.png", "图 6：前端主要功能模块")
add_para(doc, "前端使用 React + Vite 实现，为用户提供统一的科研问答工作台。")

# API table
add_heading(doc, "7. 后端 API 摘要", 1)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
hdr = table.rows[0].cells
for cell, text in zip(hdr, ["类别", "主要接口", "说明"]):
    set_cell_shading(cell, COLORS["navy"])
    set_cell_text(cell, text, True, "FFFFFF")
rows = [
    ("问答", "/chat, /chat/stream", "普通问答、流式问答、专业研究模式"),
    ("会话", "/sessions", "创建和删除对话会话"),
    ("灌入", "/ingest/rebuild, /ingest/append, /ingest/parse, /ingest/load-vec, /ingest/upload", "文档解析、重建、追加、上传和向量加载"),
    ("任务", "/tasks/{task_id}", "查询异步任务进度和结果"),
    ("知识库", "/collections", "查看和删除 Milvus collection"),
    ("技能", "/skills, /skills/template", "管理专家技能配置"),
    ("状态", "/health, /stats, /doc_summary", "健康检查、统计和文档摘要"),
    ("日志", "/logs/sessions", "查看和订阅会话级检索日志"),
]
for i, row in enumerate(rows):
    cells = table.add_row().cells
    for cell, text in zip(cells, row):
        set_cell_text(cell, text)
        if i % 2 == 0:
            set_cell_shading(cell, "F8FAFC")

add_heading(doc, "8. 项目结构摘要", 1)
add_para(doc, "pipeline/：后端 RAG 流水线、API、客户端、处理器、检索器、路由和单步任务。")
add_para(doc, "frontend/：React + Vite 前端界面，包括智能问答、知识库、专家技能、系统状态和日志页面。")
add_para(doc, "ragas_eval/：RAG 评测脚本，用于生成或执行检索与回答质量评测。")
add_para(doc, "synthetic_qa_gen/：合成问答数据生成工具，用于评测集构建和检索效果验证。")
add_para(doc, "uploads/skills/：自定义或上传的专家技能配置。")

add_heading(doc, "9. 简要总结", 1)
add_para(doc, "DP-RAG 的核心价值是把科研 PDF 转换为可检索、可引用、可追踪的知识库，并通过多策略检索和 LLM 生成能力，为用户提供带证据来源的科研文献问答体验。系统同时具备知识库管理、专家技能配置、系统状态监控、检索日志和评测工具，适合扩展为专业科研知识助手。")

doc.save(DOCX_PATH)
print(DOCX_PATH)
